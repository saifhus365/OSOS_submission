import os
from docx import Document
import fitz  # PyMuPDF

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))  # readable format

    return "\n".join(full_text)

def extract_text_from_pdf(filepath):
    doc = fitz.open(filepath)
    text_per_page = []

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        text_per_page.append((page_num, text))

    return text_per_page  # returns list of (page_number, text)

def read_document(filepath):
    ext = os.path.splitext(filepath)[-1].lower()

    if ext == ".docx":
        return [(1, extract_text_from_docx(filepath))]  # emulate single-page for docx
    elif ext == ".pdf":
        return extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type for reading: {ext}")
